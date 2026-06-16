#!/usr/bin/env python3
"""Generate bilingual handbook .docx from structured JSON input.

Input JSON format (from stdin or --input file):
{
  "title": "Vol. XX 标题",
  "subtitle": "Brain Snacks文稿校订版 · 中英双语学习手册",
  "usage_note": "S1 / S2 表示两位对谈主持人...",
  "parts": [
    {
      "title": "Part I ｜ 心态转变",
      "rows": [
        {"speaker": "S1", "english": "...", "chinese": "..."},
        ...
      ]
    },
    ...
  ],
  "vocabulary": [
    {"word": "mindset", "chinese": "心态", "freq": 7, "example": "..."},
    ...
  ],
  "idioms": [
    {"idiom": "hot potato", "chinese": "烫手山芋", "example": "...", "tip": "..."},
    ...
  ]
}

Output: .docx file
"""
import argparse, json, sys

# Color constants — Brain Snacks orange theme
THEME_PRIMARY = "E8601C"   # Brain Snacks orange (logo color)
THEME_LIGHT   = "FDE8D8"   # light orange background

# Legacy aliases (used throughout the file)
DARK_BLUE = THEME_PRIMARY
LIGHT_BLUE = THEME_LIGHT


def set_cell_shading(cell, color):
    """Set background color on a table cell via XML."""
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:shd"))
    if existing is not None:
        tc_pr.remove(existing)
    tc_pr.append(shading)


def set_run_white(run):
    """Set run font color to white."""
    from docx.shared import RGBColor
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def create_handbook(data, output_path):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # -- Styles setup --
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    # -- Title --
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(data.get("title", "Untitled"))
    run.bold = True
    run.font.size = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(data.get("subtitle", "中英双语学习手册"))
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # -- Usage note (How to Use) --
    if data.get("usage_note"):
        t = doc.add_table(rows=2, cols=1)

        # Row 0: dark blue header, white bold text
        cell0 = t.rows[0].cells[0]
        set_cell_shading(cell0, DARK_BLUE)
        cell0.text = ""
        p = cell0.paragraphs[0]
        run = p.add_run("文档结构｜How to Use This Handout")
        run.bold = True
        run.font.size = Pt(11)
        set_run_white(run)

        # Row 1: plain. Split on \n so multiple lines render properly in docx.
        cell1 = t.rows[1].cells[0]
        lines = data["usage_note"].split("\n")
        cell1.text = lines[0]
        for line in lines[1:]:
            cell1.add_paragraph(line)
        for para in cell1.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

        doc.add_paragraph()

    # -- Parts (bilingual transcript tables) --
    for part in data.get("parts", []):
        # Part title: light blue background
        t_header = doc.add_table(rows=1, cols=1)
        cell = t_header.rows[0].cells[0]
        set_cell_shading(cell, LIGHT_BLUE)
        cell.text = f"{part['title']}\nCorrected bilingual transcript / 校订双语对照"
        doc.add_paragraph()

        # Transcript table: Speaker | English | Chinese
        rows = part.get("rows", [])
        t = doc.add_table(rows=len(rows) + 1, cols=3)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row: light blue background
        headers = ["Speaker\n说话人", "English\n校订英文", "Chinese\n中文对照"]
        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            set_cell_shading(cell, LIGHT_BLUE)
            cell.text = h
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)

        # Data rows (accept both "english"/"chinese" and "en"/"zh" keys)
        for ri, row in enumerate(rows, 1):
            t.rows[ri].cells[0].text = row.get("speaker", "")
            t.rows[ri].cells[1].text = row.get("english", "") or row.get("en", "")
            t.rows[ri].cells[2].text = row.get("chinese", "") or row.get("zh", "")
            # Bold speaker name
            for para in t.rows[ri].cells[0].paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(9)
            for ci in range(1, 3):
                for paragraph in t.rows[ri].cells[ci].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

        # Column widths
        for row in t.rows:
            row.cells[0].width = Inches(0.6)
            row.cells[1].width = Inches(3.2)
            row.cells[2].width = Inches(3.2)

        doc.add_paragraph()

    # -- Vocabulary table --
    vocab = data.get("vocabulary", [])
    if vocab:
        # Section header
        t_header = doc.add_table(rows=2, cols=1)
        cell0 = t_header.rows[0].cells[0]
        set_cell_shading(cell0, DARK_BLUE)
        cell0.text = ""
        p = cell0.paragraphs[0]
        run = p.add_run("High-Frequency Academic Vocabulary / 高频学术词汇")
        run.bold = True
        run.font.size = Pt(11)
        set_run_white(run)

        t_header.rows[1].cells[0].text = "说明：以下词频基于校订后的英文稿统计，例如 mindset / mindsets 视为同一词族。"
        doc.add_paragraph()

        # Vocab data table
        t = doc.add_table(rows=len(vocab) + 1, cols=4)

        # Header: dark blue, white text
        headers = ["Word", "中文义", "Freq.", "Example sentence from the podcast / 播客原句"]
        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            set_cell_shading(cell, DARK_BLUE)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(9)
            set_run_white(run)

        for ri, v in enumerate(vocab, 1):
            t.rows[ri].cells[0].text = v.get("word", "")
            t.rows[ri].cells[1].text = v.get("chinese", "")
            t.rows[ri].cells[2].text = str(v.get("freq", ""))
            t.rows[ri].cells[3].text = v.get("example", "")
            # Bold the word
            for para in t.rows[ri].cells[0].paragraphs:
                for run in para.runs:
                    run.bold = True

        doc.add_paragraph()

    # -- Idioms table --
    idioms = data.get("idioms", [])
    if idioms:
        # Section header
        t_header = doc.add_table(rows=2, cols=1)
        cell0 = t_header.rows[0].cells[0]
        set_cell_shading(cell0, DARK_BLUE)
        cell0.text = ""
        p = cell0.paragraphs[0]
        run = p.add_run("Idioms and Natural Expressions / 地道习语与表达")
        run.bold = True
        run.font.size = Pt(11)
        set_run_white(run)

        t_header.rows[1].cells[0].text = "这些表达都来自播客原文，适合做口语积累与情境化记忆。"
        doc.add_paragraph()

        # Idiom data table
        t = doc.add_table(rows=len(idioms) + 1, cols=4)

        # Header: dark blue, white text
        headers = ["Idiom", "中文含义", "Example sentence / 原句例句", "用法提示"]
        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            set_cell_shading(cell, DARK_BLUE)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(9)
            set_run_white(run)

        for ri, idiom in enumerate(idioms, 1):
            t.rows[ri].cells[0].text = idiom.get("idiom", "")
            t.rows[ri].cells[1].text = idiom.get("chinese", "")
            t.rows[ri].cells[2].text = idiom.get("example", "")
            t.rows[ri].cells[3].text = idiom.get("tip", "")

    doc.save(output_path)
    return output_path


def create_transcript(data, output_path):
    """Generate a transcript-only .docx with speaker labels and timestamps."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(data.get("title", "Untitled") + " — 逐字稿")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()

    # Transcript entries
    for entry in data.get("transcript", []):
        p = doc.add_paragraph()

        # Timestamp + Speaker label: bold, dark blue
        ts = entry.get("timestamp", "")
        speaker = entry.get("speaker", "")
        label = f"[{ts}] {speaker}" if ts else speaker

        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xE8, 0x60, 0x1C)  # Brain Snacks orange

        # Text
        run = p.add_run(f"\n{entry.get('text', '')}")
        run.font.size = Pt(10)

        # Add spacing
        p.paragraph_format.space_after = Pt(6)

    doc.save(output_path)
    return output_path


# ── PDF generator (reportlab, no Word dependency) ────────────────

def create_handbook_pdf(data, output_path, vol_label=None):
    """Generate the bilingual handbook as a PDF using reportlab.
    Uses macOS system CJK fonts (PingFang / STHeiti fallback).
    Theme: Brain Snacks orange #E8601C / light #FDE8D8.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor, white, black
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    try:
        pdfmetrics.registerFont(TTFont("PF", "/System/Library/Fonts/PingFang.ttc", subfontIndex=0))
        pdfmetrics.registerFont(TTFont("PFB", "/System/Library/Fonts/PingFang.ttc", subfontIndex=3))
    except Exception:
        fallback = "/System/Library/Fonts/STHeiti Medium.ttc"
        pdfmetrics.registerFont(TTFont("PF", fallback))
        pdfmetrics.registerFont(TTFont("PFB", fallback))

    ORANGE = HexColor(f"#{THEME_PRIMARY}")
    LIGHT = HexColor(f"#{THEME_LIGHT}")
    GREY = HexColor("#666666")
    ROW_ALT = HexColor("#FAFAFA")

    if not vol_label:
        t = data.get("title", "")
        import re
        m = re.search(r"Vol\.?\s*\S+", t, flags=re.IGNORECASE)
        vol_label = m.group(0) if m else "Brain Snacks"

    doc = SimpleDocTemplate(
        output_path, pagesize=A4,
        leftMargin=1.8*cm, rightMargin=1.8*cm,
        topMargin=2.2*cm, bottomMargin=1.8*cm,
        title=data["title"],
    )
    st = {
        "title": ParagraphStyle("title", fontName="PFB", fontSize=20, leading=26, textColor=ORANGE, alignment=TA_CENTER, spaceAfter=6),
        "sub": ParagraphStyle("sub", fontName="PF", fontSize=11, leading=16, textColor=GREY, alignment=TA_CENTER, spaceAfter=18),
        "h2": ParagraphStyle("h2", fontName="PFB", fontSize=14, leading=20, textColor=ORANGE, spaceBefore=12, spaceAfter=6),
        "use": ParagraphStyle("use", fontName="PF", fontSize=10.5, leading=16, textColor=black, spaceAfter=14),
        "part": ParagraphStyle("part", fontName="PFB", fontSize=13, leading=18, textColor=white, alignment=TA_LEFT),
        "spk": ParagraphStyle("spk", fontName="PFB", fontSize=10, leading=14, textColor=ORANGE),
        "en": ParagraphStyle("en", fontName="PF", fontSize=10, leading=15, textColor=black),
        "zh": ParagraphStyle("zh", fontName="PF", fontSize=10, leading=15, textColor=HexColor("#333333")),
        "hdr": ParagraphStyle("hdr", fontName="PFB", fontSize=10.5, leading=14, textColor=ORANGE),
        "vt": ParagraphStyle("vt", fontName="PF", fontSize=10, leading=14, textColor=black),
        "vex": ParagraphStyle("vex", fontName="PF", fontSize=9, leading=13, textColor=GREY),
    }
    story = [
        Paragraph(data["title"], st["title"]),
        Paragraph(data["subtitle"], st["sub"]),
        Paragraph("📘 使用说明 / Usage", st["h2"]),
        Paragraph(data["usage_note"].replace("\n", "<br/>"), st["use"]),
    ]

    for part in data["parts"]:
        story.append(Spacer(1, 6))
        hdr = Table([[Paragraph(part["title"], st["part"])]], colWidths=[17.4*cm])
        hdr.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,-1), ORANGE),
            ("LEFTPADDING", (0,0), (-1,-1), 10),
            ("RIGHTPADDING", (0,0), (-1,-1), 10),
            ("TOPPADDING", (0,0), (-1,-1), 7),
            ("BOTTOMPADDING", (0,0), (-1,-1), 7),
        ]))
        story.append(hdr)

        rows = [[
            Paragraph("<b>说话人</b>", st["spk"]),
            Paragraph("<b>English</b>", st["hdr"]),
            Paragraph("<b>中文</b>", st["hdr"]),
        ]]
        for r in part["rows"]:
            rows.append([
                Paragraph(r["speaker"], st["spk"]),
                Paragraph(r.get("english", r.get("en", "")), st["en"]),
                Paragraph(r.get("chinese", r.get("zh", "")), st["zh"]),
            ])
        tbl = Table(rows, colWidths=[1.6*cm, 7.9*cm, 7.9*cm], repeatRows=1)
        tstyle = [
            ("BACKGROUND", (0,0), (-1,0), LIGHT),
            ("VALIGN", (0,0), (-1,-1), "TOP"),
            ("LEFTPADDING", (0,0), (-1,-1), 6),
            ("RIGHTPADDING", (0,0), (-1,-1), 6),
            ("TOPPADDING", (0,0), (-1,-1), 6),
            ("BOTTOMPADDING", (0,0), (-1,-1), 6),
            ("LINEBELOW", (0,0), (-1,0), 0.5, ORANGE),
            ("GRID", (0,1), (-1,-1), 0.25, HexColor("#EADBCC")),
        ]
        for i in range(2, len(rows), 2):
            tstyle.append(("BACKGROUND", (0,i), (-1,i), ROW_ALT))
        tbl.setStyle(TableStyle(tstyle))
        story.append(tbl)

    # Vocabulary
    if data.get("vocabulary"):
        story.append(Spacer(1, 10))
        story.append(Paragraph("📚 词汇表 / Vocabulary", st["h2"]))
        vrows = [[Paragraph("<b>Word</b>", st["hdr"]),
                  Paragraph("<b>中文</b>", st["hdr"]),
                  Paragraph("<b>频次</b>", st["hdr"]),
                  Paragraph("<b>原句</b>", st["hdr"])]]
        for v in data["vocabulary"]:
            vrows.append([
                Paragraph(v["word"], st["vt"]),
                Paragraph(v["chinese"], st["vt"]),
                Paragraph(str(v.get("freq", "")), st["vt"]),
                Paragraph(v.get("example", ""), st["vex"]),
            ])
        vt = Table(vrows, colWidths=[3.5*cm, 2.8*cm, 1.3*cm, 9.8*cm], repeatRows=1)
        vt.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), LIGHT),
            ("VALIGN", (0,0), (-1,-1), "TOP"),
            ("LEFTPADDING", (0,0), (-1,-1), 6),
            ("RIGHTPADDING", (0,0), (-1,-1), 6),
            ("TOPPADDING", (0,0), (-1,-1), 5),
            ("BOTTOMPADDING", (0,0), (-1,-1), 5),
            ("GRID", (0,0), (-1,-1), 0.25, HexColor("#EADBCC")),
        ]))
        story.append(vt)

    # Idioms
    if data.get("idioms"):
        story.append(Spacer(1, 14))
        story.append(Paragraph("🗣 习语表 / Idioms", st["h2"]))
        for idm in data["idioms"]:
            story.append(Paragraph(f"<b>{idm['idiom']}</b> — {idm['chinese']}", st["hdr"]))
            story.append(Paragraph(f"原句：{idm.get('example','')}", st["vex"]))
            story.append(Paragraph(f"用法：{idm.get('tip','')}", st["vt"]))
            story.append(Spacer(1, 6))

    def on_page(canvas, d):
        canvas.saveState()
        canvas.setFont("PF", 9)
        canvas.setFillColor(GREY)
        canvas.drawRightString(A4[0]-1.8*cm, A4[1]-1.2*cm, f"Brain Snacks · {vol_label}")
        canvas.drawCentredString(A4[0]/2, 1.1*cm, f"Page {d.page}")
        canvas.restoreState()

    doc.build(story, onFirstPage=on_page, onLaterPages=on_page)
    return output_path


if __name__ == "__main__":
    p = argparse.ArgumentParser()
    p.add_argument("--input", "-i", help="Handbook JSON file (default: stdin)")
    p.add_argument("--output", "-o", help="Output .docx path (optional)")
    p.add_argument("--pdf-output", help="Output .pdf path (reportlab, no Word needed)")
    p.add_argument("--transcript-input", help="Transcript JSON file for separate transcript docx")
    p.add_argument("--transcript-output", help="Output transcript .docx path")
    p.add_argument("--vol-label", help="Volume label for PDF header, e.g. 'Vol. 36'")
    args = p.parse_args()

    if not any([args.output, args.pdf_output, args.transcript_output]):
        p.error("at least one of --output / --pdf-output / --transcript-output is required")

    need_handbook = bool(args.output or args.pdf_output)
    data = None
    if need_handbook:
        if args.input:
            with open(args.input) as f:
                data = json.load(f)
        else:
            data = json.load(sys.stdin)

    if args.output and data:
        path = create_handbook(data, args.output)
        print(f"Handbook docx saved to {path}")

    if args.pdf_output and data:
        path = create_handbook_pdf(data, args.pdf_output, args.vol_label)
        print(f"Handbook PDF saved to {path}")

    if args.transcript_input and args.transcript_output:
        with open(args.transcript_input) as f:
            t_data = json.load(f)
        t_path = create_transcript(t_data, args.transcript_output)
        print(f"Transcript saved to {t_path}")
